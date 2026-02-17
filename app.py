from flask import Flask, render_template, request, jsonify
from flask_cors import CORS
import os
from groq import Groq
import PyPDF2
import docx
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

# Initialize Groq client
groq_client = Groq(api_key=os.getenv('GROQ_API_KEY'))

# Read the JD from the docx file
def read_jd():
    try:
        doc = docx.Document('Shopify Back-End Developer.docx')
        jd_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return jd_text
    except Exception as e:
        print(f"Error reading JD: {e}")
        return None

# Extract text from PDF
def extract_pdf_text(file):
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ''
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

# Extract text from DOCX
def extract_docx_text(file):
    try:
        doc = docx.Document(file)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

# Analyze resume using Groq AI
def analyze_resume_with_ai(resume_text, jd_text):
    try:
        prompt = f"""You are an expert HR recruiter analyzing resumes. Compare the following resume against the job description and provide a detailed analysis.

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}

Please provide your analysis in the following JSON format:
{{
    "overall_match_score": <number between 0-100>,
    "overall_summary": "<brief 2-3 sentence summary>",
    "strengths": ["strength 1", "strength 2", "strength 3"],
    "gaps": ["gap 1", "gap 2", "gap 3"],
    "skills_match": {{
        "matched_skills": ["skill1", "skill2", ...],
        "missing_skills": ["skill1", "skill2", ...]
    }},
    "experience_analysis": "<analysis of experience relevance>",
    "recommendation": "<HIGHLY RECOMMENDED / RECOMMENDED / MAYBE / NOT RECOMMENDED>"
}}

Be thorough and professional in your analysis."""

        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are an expert HR recruiter and resume analyzer. Always respond with valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        result = response.choices[0].message.content
        
        # Try to parse JSON from the response
        import json
        # Sometimes the AI might wrap JSON in markdown code blocks
        if "```json" in result:
            result = result.split("```json")[1].split("```")[0].strip()
        elif "```" in result:
            result = result.split("```")[1].split("```")[0].strip()
            
        analysis = json.loads(result)
        return analysis
        
    except Exception as e:
        print(f"Error in AI analysis: {e}")
        return {
            "overall_match_score": 0,
            "overall_summary": f"Error analyzing resume: {str(e)}",
            "strengths": [],
            "gaps": [],
            "skills_match": {"matched_skills": [], "missing_skills": []},
            "experience_analysis": "Analysis failed",
            "recommendation": "ERROR"
        }

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file provided'}), 400
    
    file = request.files['resume']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Extract text based on file type
    filename = file.filename.lower()
    if filename.endswith('.pdf'):
        resume_text = extract_pdf_text(file)
    elif filename.endswith('.docx'):
        resume_text = extract_docx_text(file)
    else:
        return jsonify({'error': 'Unsupported file format. Please upload PDF or DOCX'}), 400
    
    if not resume_text:
        return jsonify({'error': 'Failed to extract text from resume'}), 500
    
    # Read JD
    jd_text = read_jd()
    if not jd_text:
        return jsonify({'error': 'Failed to read job description'}), 500
    
    # Analyze with AI
    analysis = analyze_resume_with_ai(resume_text, jd_text)
    
    return jsonify(analysis)

if __name__ == '__main__':
    app.run(debug=True, port=5000)
